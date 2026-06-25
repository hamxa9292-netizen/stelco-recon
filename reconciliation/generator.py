"""
DOCX Report Generator
Produces the Debtors Reconciliation Statement Word document.
"""
import tempfile, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _fmt(n: float) -> str:
    abs_n = abs(n)
    s = f"{abs_n:,.2f}"
    return f"({s})" if n < 0 else s


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        val = kwargs.get(edge, "none")
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), val)
        if val != "none":
            tag.set(qn("w:sz"), kwargs.get(f"{edge}_sz", "4"))
            tag.set(qn("w:color"), kwargs.get(f"{edge}_color", "000000"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, align="left", red=False, size=9, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if red:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def generate_docx(location: str, result: dict, report_date: str) -> str:
    """Generate .docx and return the temp file path."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Header
    p = doc.add_paragraph()
    p.add_run("Customer Services and Billing Department").bold = True
    doc.add_paragraph("    STELCO")
    doc.add_paragraph("")
    doc.add_paragraph(f"Date: {report_date}")
    doc.add_paragraph("")

    # Title
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run(f"Debtors Reconciliation Statement - {result['location_name']}")
    r1.bold = True
    r1.underline = True
    r1.font.size = Pt(11)

    # Month/Year from report_date (format: YYYY-MM-DD or MMM YYYY)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(report_date)
    r2.bold = True
    r2.font.size = Pt(11)

    doc.add_paragraph("")

    # ── Main table ──────────────────────────────────────────────────────────
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"

    # Column widths (approx): label=9cm, elec=4cm, misc=4cm
    col_widths = [Inches(3.5), Inches(1.65), Inches(1.65)]

    # Header row
    hdr = table.add_row()
    for i, (text, w) in enumerate(zip(["", "ELECTRICITY\n(MRF)", "MISC.\n(MRF)"], col_widths)):
        c = hdr.cells[i]
        c.width = w
        _cell_text(c, text, bold=True, align="center", size=9)
        _set_cell_shading(c, "D9D9D9")
        _set_cell_border(c, bottom="single", bottom_sz="6", bottom_color="000000",
                         top="none", left="none", right="none")

    # Data rows
    for row in result["rows"]:
        is_sub   = row.get("subtotal", False)
        is_final = row.get("final", False)
        is_bold  = row.get("bold", False) or is_sub or is_final
        shade    = "EEEEEE" if is_sub else ("D9D9D9" if is_final else None)

        tr = table.add_row()
        elec_val = row["elec"]
        misc_val = row["misc"]

        cells_data = [
            (row["label"], "left",  False),
            (_fmt(elec_val) if elec_val != 0 else "0.00", "right", elec_val < 0),
            (_fmt(misc_val) if misc_val != 0 else "0.00", "right", misc_val < 0),
        ]

        for i, (text, align, red) in enumerate(cells_data):
            c = tr.cells[i]
            c.width = col_widths[i]
            _cell_text(c, text, bold=is_bold, align=align, red=red, size=9)
            if shade:
                _set_cell_shading(c, shade)

            top_border    = "single" if is_sub else "none"
            bottom_border = "double" if is_final else "none"
            _set_cell_border(
                c,
                top=top_border, top_sz="4", top_color="000000",
                bottom=bottom_border, bottom_sz="6", bottom_color="000000",
                left="none", right="none"
            )

    # Footnote
    doc.add_paragraph("")
    fn = doc.add_paragraph()
    fn_run = fn.add_run("*Credit invoice in the Total Sales")
    fn_run.italic = True
    fn_run.font.size = Pt(8)
    fn_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # ── Signature table ──────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.style = "Table Grid"
    sig_w = Inches(1.95)

    sigs = [
        ("Prepared By:", "Hamza Abdul Sattar", "Admin. Supervisor"),
        ("Checked By:",  "Ali Amir",           "Deputy Service Manager"),
        ("Approved By:", "Hussain Waheed",      "General Manager"),
    ]

    for i, (label, name, title) in enumerate(sigs):
        c = sig_table.rows[0].cells[i]
        c.width = sig_w
        p_label = c.paragraphs[0]
        r = p_label.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = "Arial"
        for _ in range(3):
            c.add_paragraph("")
        p_name = c.add_paragraph()
        rn = p_name.add_run(name)
        rn.bold = True
        rn.font.size = Pt(9)
        rn.font.name = "Arial"
        p_title = c.add_paragraph()
        rt = p_title.add_run(title)
        rt.font.size = Pt(9)
        rt.font.name = "Arial"

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name
